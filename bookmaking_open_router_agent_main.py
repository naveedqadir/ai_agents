from dotenv import load_dotenv
import os
import fitz
import asyncio
from tqdm import tqdm
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import HumanMessage
import json
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import httpx

# Load environment variables
load_dotenv()

# Get OpenRouter API Key from environment
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")

# OpenRouter API configuration
OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
HEADERS = {
    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
    "HTTP-Referer": "https://localhost:3000",  # Replace with your actual domain
    "Content-Type": "application/json"
}
# Function to extract syllabus text from PDF
def extract_syllabus_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text("text") + "\n"
    return text.strip()

# Function to split syllabus text into chapters and topics
def parse_syllabus(syllabus_text):
    """Improved syllabus parsing"""
    syllabus_lines = syllabus_text.split("\n")
    syllabus = {}
    current_chapter = None
    chapter_pattern = re.compile(r'chapter\s+(\d+)[:\s]+(.*?)(?:\s*\(.*\))?$', re.IGNORECASE)

    for line in syllabus_lines:
        line = line.strip()
        if not line:
            continue
            
        chapter_match = chapter_pattern.match(line)
        if chapter_match:
            chapter_num = chapter_match.group(1)
            chapter_title = chapter_match.group(2).strip()
            current_chapter = f"{chapter_num}. {chapter_title}"
            syllabus[current_chapter] = []
        elif current_chapter and not line.lower().startswith('chapter'):
            # Remove bullet points and clean up topic text
            clean_topic = re.sub(r'^[●\-\•]\s*', '', line)
            if clean_topic:
                syllabus[current_chapter].append(clean_topic)
    
    return syllabus

# Async function to fetch AI response
async def get_ai_response(chapter, subtopic, syllabus_context, prompt, max_retries=5):
    for attempt in range(max_retries):
        try:
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    OPENROUTER_API_URL,
                    headers=HEADERS,
                    json={
                        "model": "qwen/qwen-vl-plus:free",  # Using Qwen-v1-plus through OpenRouter
                        "messages": [
                            {"role": "user", "content": prompt}
                        ]
                    }
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
            return content.strip()
            
        except Exception as e:
            if attempt < max_retries - 1:
                wait_time = (2 ** attempt) * 2  # Exponential backoff
                print(f"\nRate limit hit. Waiting {wait_time} seconds before retry...")
                await asyncio.sleep(wait_time)
            else:
                print(f"\nFailed after {max_retries} attempts for {chapter} - {subtopic}")
                return f"Error generating content for {subtopic}. Please try again later."


def setup_document_styles(doc):
    """Setup custom styles for the document"""
    # Book Title style
    title_style = doc.styles.add_style('BookTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(50)
    title_style.paragraph_format.space_before = Pt(100)

    # TOC Title style
    toc_title_style = doc.styles.add_style('TOCTitle', WD_STYLE_TYPE.PARAGRAPH)
    toc_title_style.font.size = Pt(16)
    toc_title_style.font.bold = True
    toc_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title_style.paragraph_format.space_after = Pt(24)

    # TOC Entry styles
    toc_chapter_style = doc.styles.add_style('TOCChapter', WD_STYLE_TYPE.PARAGRAPH)
    toc_chapter_style.font.size = Pt(12)
    toc_chapter_style.font.bold = True
    toc_chapter_style.paragraph_format.space_after = Pt(6)
    toc_chapter_style.paragraph_format.left_indent = Inches(0.5)

    toc_topic_style = doc.styles.add_style('TOCTopic', WD_STYLE_TYPE.PARAGRAPH)
    toc_topic_style.font.size = Pt(11)
    toc_topic_style.paragraph_format.space_after = Pt(3)
    toc_topic_style.paragraph_format.left_indent = Inches(1.0)

    # Chapter style with better spacing
    chapter_style = doc.styles.add_style('Chapter', WD_STYLE_TYPE.PARAGRAPH)
    chapter_style.font.size = Pt(24)
    chapter_style.font.bold = True
    chapter_style.paragraph_format.page_break_before = True
    chapter_style.paragraph_format.space_after = Pt(24)
    chapter_style.paragraph_format.space_before = Pt(24)
    chapter_style.paragraph_format.keep_with_next = True
    chapter_style.font.color.rgb = RGBColor(0, 0, 0)

    # Introduction style
    intro_style = doc.styles.add_style('Introduction', WD_STYLE_TYPE.PARAGRAPH)
    intro_style.font.size = Pt(12)
    intro_style.paragraph_format.space_after = Pt(24)
    intro_style.paragraph_format.first_line_indent = Inches(0.3)
    intro_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Topic style
    topic_style = doc.styles.add_style('Topic', WD_STYLE_TYPE.PARAGRAPH)
    topic_style.font.size = Pt(16)
    topic_style.font.bold = True
    topic_style.paragraph_format.space_before = Pt(16)
    topic_style.paragraph_format.space_after = Pt(8)

    # Content style
    content_style = doc.styles.add_style('Content', WD_STYLE_TYPE.PARAGRAPH)
    content_style.font.size = Pt(11)
    content_style.paragraph_format.space_after = Pt(8)
    content_style.paragraph_format.first_line_indent = Inches(0.25)

    # Add Table of Contents style
    toc_style = doc.styles.add_style('TOC', WD_STYLE_TYPE.PARAGRAPH)
    toc_style.font.size = Pt(12)
    toc_style.paragraph_format.space_after = Pt(6)
    toc_style.paragraph_format.left_indent = Inches(0.5)

    # Add Subsection style
    subsection_style = doc.styles.add_style('Subsection', WD_STYLE_TYPE.PARAGRAPH)
    subsection_style.font.size = Pt(14)
    subsection_style.font.bold = True
    subsection_style.font.italic = True
    subsection_style.paragraph_format.space_before = Pt(12)
    subsection_style.paragraph_format.space_after = Pt(6)

def create_chapter_heading(doc, chapter_text):
    """Create properly formatted chapter heading"""
    chapter_para = doc.add_paragraph(style='Chapter')
    # Remove extra "Chapter" if it's already in the text
    if chapter_text.lower().startswith('chapter'):
        chapter_para.add_run(chapter_text)
    else:
        chapter_para.add_run(f"Chapter {chapter_text}")
    return chapter_para

async def get_chapter_intro(chapter, syllabus_context):
    """Get chapter introduction from AI"""
    prompt = f"""
    Write a brief introduction for Chapter: {chapter}
    
    Context: {syllabus_context}
    
    Follow this format EXACTLY:
    [Single paragraph introduction, no title, no headers, approximately 150 words]
    
    Requirements:
    1. Write in clear, professional tone
    2. Focus on practical relevance
    3. Provide overview of chapter content
    4. Explain importance of topics
    5. Do not use bullet points or lists
    6. Do not repeat chapter title
    7. No section headers or formatting marks
    """
    return await get_ai_response(chapter, "Introduction", syllabus_context, prompt)

async def get_topic_content(chapter, topic, syllabus_context):
    """Get detailed topic content from AI"""
    prompt = f"""
    Write educational content for: {topic}
    Chapter Context: {chapter}
    Syllabus Context: {syllabus_context}
    
    Follow this format EXACTLY:

    Note: [Single sentence overview of the topic]

    1.1 [First Aspect of {topic}]
    [Detailed explanation in paragraph form]

    1.2 [Second Aspect of {topic}]
    [Detailed explanation in paragraph form]

    1.3 [Third Aspect of {topic}]
    [Detailed explanation in paragraph form]

    Requirements:
    1. No topic title repetition
    2. No bullet points or asterisks
    3. Write paragraphs, not lists
    4. Keep subtopic titles brief
    5. Focus on specific topic content
    6. No introductory or concluding remarks
    7. No formatting marks or special characters
    """
    return await get_ai_response(chapter, topic, syllabus_context, prompt)

async def get_chapter_review(chapter, topics, syllabus_context):
    """Get chapter review questions"""
    topics_list = "\n".join(f"- {topic}" for topic in topics)
    prompt = f"""
    Create review questions for Chapter: {chapter}
    
    Topics to cover:
    {topics_list}
    
    Follow this format EXACTLY:
    Review Questions for {chapter}

    Topic: [First Topic Name]
    1. [Conceptual question about the topic]
    2. [Practical application question]
    3. [Problem-solving scenario question]

    Topic: [Second Topic Name]
    4. [Conceptual question about the topic]
    5. [Practical application question]
    6. [Problem-solving scenario question]

    Requirements:
    1. Number questions sequentially across all topics
    2. No bullet points or asterisks
    3. No section headers in special formatting
    4. Keep questions clear and focused
    5. Mix question types for each topic
    6. Include practical scenarios
    7. Write in professional tone
    """
    return await get_ai_response(chapter, "Review Questions", syllabus_context, prompt)

async def ai_parse_syllabus(syllabus_text):
    """Use AI to parse syllabus into structured format"""
    prompt = f"""
    You are a syllabus parsing assistant. Parse this syllabus into chapters and topics.

    Rules:
    1. Only include main chapters and their direct topics
    2. Remove any 'Topics:' headers
    3. Ignore 'The individual shall be able to:' sections
    4. Clean up bullet points and numbering
    5. Keep actual educational topics only

    Return ONLY valid JSON in this exact format:
    {{
        "1. Work Organization and Management": [
            "Creativity in the design of circuits",
            "Critical thinking in circuit design"
        ]
    }}

    Syllabus Text:
    {syllabus_text}
    """
    
    try:
        # First try to get structured response
        response = await get_ai_response("Syllabus Parsing", "Structure", syllabus_text, prompt)
        
        # Clean the response to ensure it only contains the JSON part
        json_str = re.search(r'\{.*\}', response.replace('\n', ' '), re.DOTALL)
        if not json_str:
            raise ValueError("No JSON found in response")
            
        cleaned_json = json_str.group(0)
        parsed_syllabus = json.loads(cleaned_json)
        
        # Validate the structure
        if not isinstance(parsed_syllabus, dict):
            raise ValueError("Invalid syllabus structure")
            
        # Ensure all entries are properly formatted
        validated_syllabus = {}
        for chapter, topics in parsed_syllabus.items():
            if not isinstance(topics, list):
                continue
                
            # Clean chapter name
            clean_chapter = chapter.strip()
            if not clean_chapter.lower().startswith('chapter'):
                clean_chapter = f"Chapter {clean_chapter}"
                
            # Clean topics
            clean_topics = [
                topic.strip() for topic in topics 
                if topic.strip() and not topic.lower().startswith('topic')
            ]
            
            if clean_topics:
                validated_syllabus[clean_chapter] = clean_topics
        
        if not validated_syllabus:
            raise ValueError("No valid chapters found")
            
        return validated_syllabus
        
    except Exception as e:
        print(f"AI parsing failed: {str(e)}")
        print("Falling back to regular parsing...")
        return parse_syllabus(syllabus_text)

def parse_content_sections(content):
    """Parse content into properly formatted sections"""
    sections = []
    current_section = None
    current_topic = None
    
    for line in content.split('\n'):
        line = line.rstrip()
        
        if not line:
            continue
            
        # Handle review questions header
        if line.startswith('Review Questions'):
            current_section = {'type': 'review_header', 'content': line}
            sections.append(current_section)
            continue
            
        # Handle topic headers in review questions
        if line.startswith('Topic:'):
            current_topic = {'type': 'review_topic', 'content': line.replace('Topic:', '').strip()}
            sections.append(current_topic)
            continue
            
        # Handle numbered questions
        if re.match(r'^\d+\.', line):
            current_section = {'type': 'review_question', 'content': line}
            sections.append(current_section)
            continue
            
        # Handle other content types as before
        if line.startswith('Chapter'):
            current_section = {'type': 'chapter', 'content': line}
            sections.append(current_section)
            
        elif line.startswith('Topic:'):
            continue  # Skip topic headers as they're handled separately
            
        elif line.startswith('Note:'):
            current_section = {'type': 'note', 'content': line.replace('Note:', '').strip()}
            sections.append(current_section)
            
        elif re.match(r'^\d+\.\d+\s+', line):
            current_section = {'type': 'subtopic', 'content': line}
            sections.append(current_section)
            
        elif line.lstrip().startswith(('•', '*', '-')):
            item = re.sub(r'^[•\*\-]\s*', '', line.lstrip())
            if item:
                sections.append({'type': 'bullet_list', 'items': [item]})
                
        else:
            current_section = {'type': 'paragraph', 'content': line}
            sections.append(current_section)
    
    return sections

def add_formatted_content(doc, content):
    """Add properly formatted content to document"""
    sections = parse_content_sections(content)
    
    for section in sections:
        if section['type'] == 'chapter':
            para = doc.add_paragraph(style='Chapter')
            para.add_run(section['content'])
            
        elif section['type'] == 'topic_header':
            para = doc.add_paragraph(style='Topic')
            para.add_run(section['content'])
            
        elif section['type'] == 'note':
            para = doc.add_paragraph(style='Subsection')
            para.add_run("Note: " + section['content'])
            
        elif section['type'] == 'subtopic':
            para = doc.add_paragraph(style='Topic')
            para.add_run(section['content'])
            
        elif section['type'] == 'bullet_list':
            for item in section['items']:
                para = doc.add_paragraph(style='Content')
                para.add_run("• " + item)
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.first_line_indent = Inches(-0.25)
            
        elif section['type'] == 'paragraph':
            para = doc.add_paragraph(style='Content')
            para.add_run(section['content'])
            para.paragraph_format.space_after = Pt(12)
            
        elif section['type'] == 'review_header':
            para = doc.add_paragraph(style='Topic')
            para.add_run(section['content'])
            para.paragraph_format.space_before = Pt(24)
            
        elif section['type'] == 'review_topic':
            para = doc.add_paragraph(style='Subsection')
            para.add_run(section['content'])
            para.paragraph_format.space_before = Pt(12)
            
        elif section['type'] == 'review_question':
            para = doc.add_paragraph(style='Content')
            para.add_run(section['content'])
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)

# Async function to process syllabus and generate Word document
async def process_syllabus(pdf_path, output_filename="Generated_Book.docx"):
    syllabus_text = extract_syllabus_from_pdf(pdf_path)
    
    # Use AI parsing instead of regular parsing
    syllabus = await ai_parse_syllabus(syllabus_text)
    
    # Create new document with proper styles
    doc = Document()
    setup_document_styles(doc)

    # Extract subject name from PDF filename
    subject_name = os.path.splitext(os.path.basename(pdf_path))[0].replace('-', ' ').title()
    
    # Create title page
    title = doc.add_paragraph(style='BookTitle')
    title.add_run(subject_name)
    doc.add_page_break()

    # Add Table of Contents
    toc_title = doc.add_paragraph("Table of Contents", style='TOCTitle')
    
    for chapter, topics in syllabus.items():
        toc_chapter = doc.add_paragraph(style='TOCChapter')
        toc_chapter.add_run(chapter)
        for topic in topics:
            toc_topic = doc.add_paragraph(style='TOCTopic')
            toc_topic.add_run(topic)

    doc.add_page_break()

    current_chapter = None
    for chapter, subtopics in tqdm(syllabus.items(), desc="Processing Chapters"):
        chapter_number = chapter.split(':')[0].replace('Chapter', '').strip()
        
        if current_chapter != chapter_number:
            current_chapter = chapter_number
            create_chapter_heading(doc, chapter)

        # # Add chapter introduction with new style
        intro = await get_chapter_intro(chapter, syllabus_text)
        add_formatted_content(doc, intro)

        for subtopic in tqdm(subtopics, desc=f"Processing topics", leave=False):
            print(subtopic)
            if subtopic.strip() == "":
                continue

            # Add topic heading only once
            topic_para = doc.add_paragraph(style='Topic')
            topic_para.add_run(subtopic)

            # Get and format content without repeating the topic
            content = await get_topic_content(chapter, subtopic, syllabus_text)
            add_formatted_content(doc, content)

            # Save progress after each topic
            doc.save(output_filename)
            await asyncio.sleep(2)

        # Add review questions section
        doc.add_paragraph()  # Add extra space
        review_heading = doc.add_paragraph("Review Questions", style='Topic')
        questions = await get_chapter_review(chapter, subtopics, syllabus_text)
        add_formatted_content(doc, questions)

        # Add page break after chapter
        doc.add_page_break()
        doc.save(output_filename)

    print(f"\nDocument saved as {output_filename}")

# Run async processing
if __name__ == "__main__":
    pdf_path = "Electronics-Sylabuss.pdf"  # Replace with your syllabus PDF path
    asyncio.run(process_syllabus(pdf_path))
