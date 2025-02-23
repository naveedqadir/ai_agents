from dotenv import load_dotenv
import os
import fitz
import asyncio
from tqdm import tqdm
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import httpx
import json
import re
from typing import Dict, List
import logging
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Configuration
class Config:
    OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
    OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
    HEADERS = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "HTTP-Referer": "http://localhost:3000",
        "Content-Type": "application/json"
    }
    MAX_RETRIES = 5
    MODEL = "qwen/qwen-vl-plus:free"

class APIError(Exception):
    """Custom exception for API-related errors."""
    pass

def extract_syllabus_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF syllabus file."""
    try:
        doc = fitz.open(pdf_path)
        text = "\n".join(page.get_text("text").strip() for page in doc)
        doc.close()
        return text
    except Exception as e:
        logger.error(f"Failed to extract PDF text: {e}")
        raise

async def api_request(prompt: str, max_retries: int = Config.MAX_RETRIES) -> str:
    """Make async API request with exponential backoff."""
    payload = {"model": Config.MODEL, "messages": [{"role": "user", "content": prompt}]}
    async with httpx.AsyncClient(timeout=30.0) as client:
        for attempt in range(max_retries):
            try:
                response = await client.post(Config.OPENROUTER_API_URL, headers=Config.HEADERS, json=payload)
                response.raise_for_status()
                return response.json()['choices'][0]['message']['content'].strip()
            except (httpx.HTTPError, KeyError) as e:
                if attempt == max_retries - 1:
                    raise APIError(f"API request failed after {max_retries} attempts: {e}")
                wait_time = (2 ** attempt) * 2
                logger.warning(f"API request failed. Retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)

async def ai_parse_syllabus(syllabus_text: str) -> Dict[str, List[str]]:
    """Parse syllabus into chapters and topics using AI."""
    prompt = f"""
    Parse the following syllabus text into a structured format with chapters and their topics. Adapt to varying formats (e.g., 'Chapter 1: Title', 'Unit I - Title', numbered sections, bullet points).

    Syllabus Text:
    {syllabus_text}

    Format Requirements:
    - Return ONLY valid JSON:
      {{
        "Chapter 1: [Title]": ["Topic 1", "Topic 2", ...],
        "Chapter 2: [Title]": ["Topic 1", "Topic 2", ...],
        ...
      }}
    - Use "Chapter [Number]: [Title]" as keys
    - List topics as plain strings without numbering or bullets

    Content Rules:
    - Identify chapter/section/unit headings and their titles
    - Extract associated topics/subtopics
    - Ignore preamble, objectives, or metadata
    - Remove bullet points, numbers, or formatting markers
    - Handle inconsistent formatting
    - Skip empty or irrelevant lines
    - Ensure each chapter has at least one topic
    """
    try:
        response = await api_request(prompt)
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if not json_match:
            raise ValueError("No valid JSON found in AI response")
        syllabus = json.loads(json_match.group(0))
        return {k: v for k, v in syllabus.items() if isinstance(v, list) and v}
    except Exception as e:
        logger.error(f"AI parsing failed: {e}")
        raise APIError(f"Failed to parse syllabus with AI: {e}")

def setup_document_styles(doc: Document) -> None:
    """Set up book-like styles for the Word document."""
    styles_config = {
        'BookTitle': {'size': 28, 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER, 'space_after': 50},
        'TOCTitle': {'size': 16, 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER, 'space_after': 12},
        'TOCEntry': {'size': 12, 'space_after': 4, 'left_indent': 0.5},
        'Chapter': {'size': 20, 'bold': True, 'space_before': 36, 'space_after': 18, 'page_break': True},
        'Topic': {'size': 16, 'bold': True, 'space_before': 18, 'space_after': 8},
        'Subsection': {'size': 14, 'bold': True, 'italic': True, 'space_before': 12, 'space_after': 6},
        'Content': {'size': 11, 'space_after': 6, 'first_line_indent': 0.25}
    }

    for name, config in styles_config.items():
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(config['size'])
        if config.get('bold'):
            style.font.bold = True
        if config.get('italic'):
            style.font.italic = True
        if 'align' in config:
            style.paragraph_format.alignment = config['align']
        if 'space_before' in config:
            style.paragraph_format.space_before = Pt(config['space_before'])
        if 'space_after' in config:
            style.paragraph_format.space_after = Pt(config['space_after'])
        if 'left_indent' in config:
            style.paragraph_format.left_indent = Inches(config['left_indent'])
        if 'first_line_indent' in config:
            style.paragraph_format.first_line_indent = Inches(config['first_line_indent'])
        if 'page_break' in config:
            style.paragraph_format.page_break_before = True

async def get_chapter_intro(chapter: str, syllabus_context: str, doc: Document, output_filename: str) -> str:
    """Generate and save chapter introduction with clean formatting."""
    prompt = f"""
    Generate a polished introduction for '{chapter}' for a professional book.

    Context: {syllabus_context}

    Format:
    - Plain text, no markdown, hashtags, numbering, or special characters
    - Single paragraph, 150-200 words, styled as 'Content' (11pt, justified, 0.25in indent)

    Content:
    - Formal, academic tone
    - Overview of chapter topics
    - Emphasize practical relevance and foundational importance
    - Avoid repeating chapter title
    - No lists, bullet points, or subsection headers
    """
    content = await api_request(prompt)
    cleaned_content = re.sub(r'[#*\-]+\s*', '', content)  # Remove any residual markers
    add_formatted_content(doc, cleaned_content)
    doc.save(output_filename)
    logger.info(f"Saved progress after chapter intro: {chapter}")
    return cleaned_content

async def get_topic_content(chapter: str, topic: str, syllabus_context: str, doc: Document, output_filename: str) -> str:
    """Generate and save topic content with clean formatting."""
    prompt = f"""
    Generate educational content for '{topic}' in '{chapter}' for a professional book.

    Context: {syllabus_context}

    Format:
    - Plain text, no markdown, hashtags, numbering, or special characters
    - Structure:
      - 'Note: [One-sentence overview]' (for 'Subsection' style: 14pt, bold, italic)
      - 3 subsections, each with:
        - '[Brief Title]' (no numbers like '1.1', for 'Subsection' style: 14pt, bold, italic)
        - Paragraph (100-150 words, for 'Content' style: 11pt, justified, 0.25in indent)

    Content:
    - Academic tone, technical precision
    - Distinct aspects (e.g., concepts, applications, analysis)
    - Practical examples
    - No lists or introductory remarks beyond note
    """
    content = await api_request(prompt)
    cleaned_content = re.sub(r'[#*\-]+\s*|\[\d+\.\d+\s*', '[', content)  # Remove markers and numbers before titles
    add_formatted_content(doc, cleaned_content)
    doc.save(output_filename)
    logger.info(f"Saved progress after topic content: {topic}")
    return cleaned_content

async def get_chapter_review(chapter: str, topics: List[str], syllabus_context: str, doc: Document, output_filename: str) -> str:
    """Generate and save chapter review questions with clean formatting."""
    topics_list = "\n".join(topics)
    prompt = f"""
    Generate review questions for '{chapter}' for a professional book.

    Topics:
    {topics_list}

    Format:
    - Plain text, no markdown, hashtags, or special characters except question numbers
    - Structure:
      - 'Review Questions' (for 'Topic' style: 16pt, bold)
      - For each topic:
        - 'Topic: [Name]' (for 'Subsection' style: 14pt, bold, italic)
        - 3 numbered questions (e.g., '1. Text', for 'Content' style: 11pt, 0.25in indent)

    Content:
    - Academic tone
    - Per topic: 1 conceptual, 1 practical, 1 problem-solving question
    - Number consecutively across topics
    """
    content = await api_request(prompt)
    cleaned_content = re.sub(r'[#*\-]+\s*', '', content)  # Remove unwanted markers, keep question numbers
    add_formatted_content(doc, cleaned_content)
    doc.save(output_filename)
    logger.info(f"Saved progress after chapter review: {chapter}")
    return cleaned_content

def add_formatted_content(doc: Document, content: str) -> None:
    """Add content to document with book-appropriate formatting."""
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('Review Questions'):
            doc.add_paragraph(line, style='Topic')
        elif line.startswith('Topic:') or line.startswith('Note:') or (not re.match(r'^\d+\.\s', line) and ':' not in line and len(line.split()) <= 5):
            doc.add_paragraph(line, style='Subsection')
        elif re.match(r'^\d+\.\s', line):
            para = doc.add_paragraph(line, style='Content')
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)
        else:
            doc.add_paragraph(line, style='Content')

async def process_syllabus(pdf_path: str, output_filename: str = "Generated_Book.docx") -> None:
    """Generate a properly formatted book with real-time saving and clean formatting."""
    try:
        # Initialization
        syllabus_text = extract_syllabus_from_pdf(pdf_path)
        syllabus = await ai_parse_syllabus(syllabus_text)
        doc = Document()
        setup_document_styles(doc)

        # Title Page
        subject_name = Path(pdf_path).stem.replace('-', ' ').title()
        doc.add_paragraph(f"{subject_name}", style='BookTitle')
        doc.add_paragraph("A Comprehensive Study Guide", style='TOCTitle')
        doc.add_page_break()

        # Table of Contents
        doc.add_paragraph("Table of Contents", style='TOCTitle')
        for chapter, topics in syllabus.items():
            doc.add_paragraph(chapter, style='TOCEntry')
            for topic in topics:
                para = doc.add_paragraph(topic, style='TOCEntry')
                para.paragraph_format.left_indent = Inches(0.75)
        doc.add_page_break()
        doc.save(output_filename)
        logger.info("Saved initial document structure")

        # Chapters
        for chapter, topics in tqdm(syllabus.items(), desc="Chapters"):
            doc.add_paragraph(chapter, style='Chapter')
            
            await get_chapter_intro(chapter, syllabus_text, doc, output_filename)
            
            for topic in tqdm(topics, desc=f"Topics", leave=False):
                if not topic.strip():
                    continue
                doc.add_paragraph(topic, style='Topic')
                await get_topic_content(chapter, topic, syllabus_text, doc, output_filename)
            
            await get_chapter_review(chapter, topics, syllabus_text, doc, output_filename)
            doc.add_page_break()

        doc.save(output_filename)
        logger.info(f"Book fully generated: {output_filename}")

    except Exception as e:
        logger.error(f"Book generation failed: {e}")
        doc.save(output_filename)
        raise

if __name__ == "__main__":
    pdf_path = "Electronics-Sylabuss.pdf"
    asyncio.run(process_syllabus(pdf_path))